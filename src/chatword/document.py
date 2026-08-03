"""Document inspection and text extraction helpers."""

from __future__ import annotations

from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Any

WORD_EXTENSIONS = {".docx"}
LEGACY_WORD_EXTENSIONS = {".doc"}
PDF_EXTENSIONS = {".pdf"}


class ChatWordError(Exception):
    """Base exception for ChatWord document operations."""


class MissingDependencyError(ChatWordError):
    """Raised when an optional parser dependency is not installed."""

    def __init__(self, package: str, extra: str) -> None:
        self.package = package
        self.extra = extra
        super().__init__(
            f"Missing optional dependency {package!r}. Install with "
            f"`pip install ChatWord[{extra}]`."
        )


class UnsupportedDocumentError(ChatWordError):
    """Raised when ChatWord does not know how to parse a file."""


class DocumentParseError(ChatWordError):
    """Raised when a parser is available but cannot read a document."""


@dataclass(frozen=True)
class DocumentInfo:
    """Basic document metadata that is safe to render in CLI output."""

    path: str
    kind: str
    suffix: str
    size_bytes: int
    parser: str = "none"
    details: dict[str, Any] = field(default_factory=dict)
    warning: str | None = None

    def to_dict(self) -> dict[str, Any]:
        """Return a JSON-serializable representation."""

        return asdict(self)


def detect_document_kind(path: str | Path) -> str:
    """Detect a supported document kind from a path suffix."""

    suffix = Path(path).suffix.lower()
    if suffix in WORD_EXTENSIONS:
        return "docx"
    if suffix in LEGACY_WORD_EXTENSIONS:
        return "doc"
    if suffix in PDF_EXTENSIONS:
        return "pdf"
    return "unknown"


def _resolve_document_file(path: str | Path) -> Path:
    resolved = Path(path).expanduser().resolve()
    if not resolved.exists():
        raise DocumentParseError(f"Document not found: {resolved}")
    if not resolved.is_file():
        raise DocumentParseError(f"Document path is not a file: {resolved}")
    return resolved


def inspect_document(path: str | Path) -> DocumentInfo:
    """Inspect a document without extracting its full text."""

    resolved = _resolve_document_file(path)
    suffix = resolved.suffix.lower()
    kind = detect_document_kind(resolved)
    size_bytes = resolved.stat().st_size

    if kind == "docx":
        return _inspect_docx(resolved, suffix, size_bytes)
    if kind == "pdf":
        return _inspect_pdf(resolved, suffix, size_bytes)
    if kind == "doc":
        return DocumentInfo(
            path=str(resolved),
            kind=kind,
            suffix=suffix,
            size_bytes=size_bytes,
            warning="Legacy .doc parsing needs an external converter such as LibreOffice.",
        )
    return DocumentInfo(
        path=str(resolved),
        kind=kind,
        suffix=suffix,
        size_bytes=size_bytes,
        warning="Unsupported document type.",
    )


def extract_text(path: str | Path) -> str:
    """Extract plain text from a supported document."""

    resolved = _resolve_document_file(path)
    kind = detect_document_kind(resolved)
    if kind == "docx":
        return _extract_docx_text(resolved)
    if kind == "pdf":
        return _extract_pdf_text(resolved)
    if kind == "doc":
        raise UnsupportedDocumentError(
            "Legacy .doc extraction needs conversion to .docx first."
        )
    raise UnsupportedDocumentError(f"Unsupported document type: {resolved.suffix or 'none'}")


def _inspect_docx(path: Path, suffix: str, size_bytes: int) -> DocumentInfo:
    try:
        from docx import Document  # type: ignore[reportMissingImports]
    except ImportError as exc:  # pragma: no cover - depends on optional env
        error = MissingDependencyError("python-docx", "word")
        return DocumentInfo(
            path=str(path),
            kind="docx",
            suffix=suffix,
            size_bytes=size_bytes,
            parser="missing",
            warning=str(error),
        )

    try:
        document = Document(path)
        props = document.core_properties
        details = {
            "paragraphs": len(document.paragraphs),
            "tables": len(document.tables),
            "sections": len(document.sections),
            "title": props.title or None,
            "author": props.author or None,
            "created": _format_datetime(props.created),
            "modified": _format_datetime(props.modified),
        }
    except Exception as exc:
        raise DocumentParseError(f"Could not inspect DOCX file: {exc}") from exc
    return DocumentInfo(
        path=str(path),
        kind="docx",
        suffix=suffix,
        size_bytes=size_bytes,
        parser="python-docx",
        details=details,
    )


def _inspect_pdf(path: Path, suffix: str, size_bytes: int) -> DocumentInfo:
    try:
        from pypdf import PdfReader  # type: ignore[reportMissingImports]
    except ImportError as exc:  # pragma: no cover - depends on optional env
        error = MissingDependencyError("pypdf", "pdf")
        return DocumentInfo(
            path=str(path),
            kind="pdf",
            suffix=suffix,
            size_bytes=size_bytes,
            parser="missing",
            warning=str(error),
        )

    try:
        reader = PdfReader(str(path))
        encrypted = bool(reader.is_encrypted)
        if encrypted:
            return DocumentInfo(
                path=str(path),
                kind="pdf",
                suffix=suffix,
                size_bytes=size_bytes,
                parser="pypdf",
                details={"encrypted": True},
                warning="Encrypted PDF content requires a password before inspection.",
            )
        metadata = reader.metadata or {}
        details = {
            "pages": len(reader.pages),
            "encrypted": False,
            "metadata": _clean_pdf_metadata(metadata),
        }
    except Exception as exc:
        raise DocumentParseError(f"Could not inspect PDF file: {exc}") from exc
    return DocumentInfo(
        path=str(path),
        kind="pdf",
        suffix=suffix,
        size_bytes=size_bytes,
        parser="pypdf",
        details=details,
    )


def _extract_docx_text(path: Path) -> str:
    try:
        from docx import Document  # type: ignore[reportMissingImports]
    except ImportError as exc:
        raise MissingDependencyError("python-docx", "word") from exc

    try:
        document = Document(path)
        parts: list[str] = []
        parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append("\t".join(cells))
    except Exception as exc:
        raise DocumentParseError(f"Could not extract DOCX text: {exc}") from exc
    return "\n".join(parts).strip() + "\n"


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader  # type: ignore[reportMissingImports]
    except ImportError as exc:
        raise MissingDependencyError("pypdf", "pdf") from exc

    try:
        reader = PdfReader(str(path))
        if reader.is_encrypted:
            raise DocumentParseError("Encrypted PDF content requires a password before text extraction.")
        pages = []
        for index, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                pages.append(f"--- page {index} ---\n{text.strip()}")
    except DocumentParseError:
        raise
    except Exception as exc:
        raise DocumentParseError(f"Could not extract PDF text: {exc}") from exc
    return "\n\n".join(pages).strip() + "\n"


def _format_datetime(value: Any) -> str | None:
    if value is None:
        return None
    isoformat = getattr(value, "isoformat", None)
    return str(isoformat()) if callable(isoformat) else str(value)


def _clean_pdf_metadata(metadata: Any) -> dict[str, str]:
    cleaned: dict[str, str] = {}
    for key, value in dict(metadata).items():
        if value is not None:
            cleaned[str(key).lstrip("/")] = str(value)
    return cleaned
